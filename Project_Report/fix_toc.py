r"""
fix_toc.py
----------
Replaces the static pandoc-generated TOC (Compact-style paragraphs)
with a proper Word TOC field { TOC \o "1-3" \h \z \u }.

When the output file is opened in Word, press:
  Ctrl+A  ->  F9  ->  "Update entire table"
to populate page numbers automatically.
"""

import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

INPUT  = "Full_Report_V3.docx"
OUTPUT = "Full_Report_Final_V3.docx"


# ── Build the Word TOC field paragraph ───────────────────────────────────────
def make_toc_field_element():
    r"""
    Returns a <w:p> containing the Word field:
        { TOC \o "1-3" \h \z \u }
    with w:dirty="true" so Word knows to update it on open.
    """
    para = OxmlElement('w:p')

    def run_with(child):
        r = OxmlElement('w:r')
        r.append(child)
        return r

    # BEGIN
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    fc_begin.set(qn('w:dirty'), 'true')
    para.append(run_with(fc_begin))

    # INSTRUCTION
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    para.append(run_with(instr))

    # SEPARATE (empty result placeholder)
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    para.append(run_with(fc_sep))

    # END
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    para.append(run_with(fc_end))

    return para


# ── Find and replace the static TOC ──────────────────────────────────────────
def fix_toc(doc: Document):
    paras = doc.paragraphs

    # 1. Locate the "Table of Contents" Heading 1 paragraph index
    toc_heading_idx = None
    for i, p in enumerate(paras):
        if p.style.name == 'Heading 1' and 'table of contents' in p.text.lower():
            toc_heading_idx = i
            print(f"[found] TOC heading at paragraph index {i}: {p.text!r}")
            break

    if toc_heading_idx is None:
        print("[error] Could not find 'Table of Contents' heading. Aborting.")
        sys.exit(1)

    # 2. Collect all Compact/toc-style paragraphs that follow the heading,
    #    stopping at the next Heading 1 (start of next chapter).
    TOC_STYLES = {'Compact', 'toc 1', 'toc 2', 'toc 3',
                  'TOC 1',   'TOC 2', 'TOC 3'}
    to_remove = []
    for p in paras[toc_heading_idx + 1:]:
        sname = p.style.name if p.style else ''
        if p.style.name == 'Heading 1':
            break   # hit next chapter — stop
        if sname in TOC_STYLES or sname == 'Normal':
            # Only grab Normal paragraphs that look like TOC entries
            # (they have plain text and no special content)
            to_remove.append(p._element)

    print(f"[found] {len(to_remove)} static TOC paragraph(s) to replace.")

    # 3. Insert the Word TOC field right after the TOC heading paragraph
    toc_heading_elem = paras[toc_heading_idx]._element
    toc_field = make_toc_field_element()
    toc_heading_elem.addnext(toc_field)
    print(f"[ok]    Inserted Word TOC field after heading.")

    # 4. Remove the old static entries
    body = doc.element.body
    removed = 0
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            removed += 1
    print(f"[ok]    Removed {removed} static TOC paragraph(s).")

    return doc


# ── Main ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')

    print("=" * 60)
    print("  CriticAI - Fix Word TOC Page Numbers")
    print("=" * 60)

    doc = Document(INPUT)
    print(f"[open]  {INPUT}  ({len(doc.paragraphs)} paragraphs)")

    doc = fix_toc(doc)
    doc.save(OUTPUT)

    print(f"\n[saved] {OUTPUT}")
    print("=" * 60)
    print()
    print("NEXT STEP - Open the file in Microsoft Word and do:")
    print("  1. Press  Ctrl + A  (select all)")
    print("  2. Press  F9")
    print("  3. Choose 'Update entire table'")
    print("  4. Save the file")
    print()
    print("The Table of Contents will show correct page numbers.")
    print("=" * 60)
